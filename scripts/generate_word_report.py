"""
Generates API_CONNECTIONS_MAP.docx from the markdown file.
Run: py -3 scripts/generate_word_report.py
"""

import sys
sys.path.insert(0, r'C:\Users\User\AppData\Roaming\Python\Python313\site-packages')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_paragraph(doc, text, bold=False, size=11, color=None,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6,
                  italic=False, rtl=False):
    p = doc.add_paragraph()
    p.alignment = align
    pPr = p._p.get_or_add_pPr()
    if rtl:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_heading(doc, text, level=1):
    colors = {1: '1B3A6B', 2: '2E5FA3', 3: '3A7EBF'}
    sizes  = {1: 18, 2: 15, 3: 13}
    p = add_paragraph(doc, text, bold=True, size=sizes.get(level, 13),
                      color=colors.get(level, '1B3A6B'),
                      space_before=12, space_after=6)
    # bottom border for h1
    if level == 1:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2E5FA3')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_table(doc, headers, rows, col_widths=None, header_bg='1B3A6B'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def add_code_block(doc, text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        # light grey background on paragraph level via shading the run
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()

def add_info_box(doc, rows_kv, title=None, bg='EBF2FF'):
    """Two-column key-value box."""
    if title:
        add_paragraph(doc, title, bold=True, size=10, color='1B3A6B',
                      space_before=4, space_after=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for k, v in rows_kv:
        row = table.add_row()
        # key
        kc = row.cells[0]
        set_cell_bg(kc, '1B3A6B')
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True
        kr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        kr.font.size = Pt(9)
        kc.width = Inches(1.8)
        # value
        vc = row.cells[1]
        set_cell_bg(vc, bg)
        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.size = Pt(9)
    doc.add_paragraph()

# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# page margins
sec = doc.sections[0]
sec.top_margin    = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)

# default font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ── COVER PAGE ────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = cover_title.add_run('خريطة اتصالات API')
rt.bold = True
rt.font.size = Pt(26)
rt.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('API Connections Map')
rs.bold = True
rs.font.size = Pt(20)
rs.font.color.rgb = RGBColor(0x2E, 0x5F, 0xA3)

doc.add_paragraph()

sysname = doc.add_paragraph()
sysname.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = sysname.add_run('Tatubu School Management System')
rn.font.size = Pt(14)
rn.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('تاريخ التقرير / Date',       'فبراير 2026 / February 2026'),
    ('النظام / System',             'Tatubu School Management System'),
    ('المنطقة / Region',            'سلطنة عُمان / Sultanate of Oman'),
    ('التصنيف / Classification',    'سري / Confidential'),
]
for i, (k, v) in enumerate(meta_data):
    r = meta_table.rows[i]
    set_cell_bg(r.cells[0], '1B3A6B')
    set_cell_bg(r.cells[1], 'EBF2FF')
    kp = r.cells[0].paragraphs[0]
    kr = kp.add_run(k)
    kr.bold = True; kr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); kr.font.size = Pt(9)
    vp = r.cells[1].paragraphs[0]
    vr = vp.add_run(v)
    vr.font.size = Pt(9)
    r.cells[0].width = Inches(2.5)
    r.cells[1].width = Inches(3.5)

doc.add_page_break()

# ── SECTION 1: Overview ───────────────────────────────────────────────────────

add_heading(doc, '1. نظرة عامة / System Overview', 1)
add_paragraph(doc,
    'يوضح هذا التقرير جميع الاتصالات البرمجية (API Connections) في نظام تطبوبي لإدارة المدارس، '
    'مع تفصيل كل اتصال من حيث المزود، والبروتوكول، وطريقة المصادقة، والبيانات المنقولة.',
    size=10, rtl=True)
add_paragraph(doc,
    'This report documents all API connections within the Tatubu School Management System, '
    'detailing each connection\'s provider, protocol, authentication method, and data transferred.',
    size=10)

doc.add_paragraph()

# ── SECTION 2: Full Connections Table ────────────────────────────────────────

add_heading(doc, '2. جدول الاتصالات الكامل / Complete Connections Table', 1)

headers = ['#', 'المُرسِل\nSender', 'المستقبِل\nReceiver', 'المزود\nProvider',
           'البروتوكول\nProtocol', 'المصادقة\nAuth', 'الاتجاه\nDirection']
rows = [
    ['1', 'React Frontend\n(Browser)',       'Flask Backend API',          'Self-hosted',            'HTTPS / REST',          'JWT Bearer Token',         'ثنائي / Bidirectional'],
    ['2', 'Flask Backend',                   'MySQL Database',              'Self-hosted',            'TCP :3306 (PyMySQL)',    'Username + Password',      'ثنائي / Bidirectional'],
    ['3', 'Flask Backend\n(Flask-Limiter)',   'Redis Cache',                 'Self-hosted',            'TCP :6379 (RESP)',       'Password (optional)',      'ثنائي / Bidirectional'],
    ['4', 'Flask Backend\n(Notification Svc)','Evolution API',              'Self-hosted (Docker)',   'HTTP / REST',           'API Key Header',           'أحادي / Outbound only'],
    ['5', 'Evolution API\n(Baileys lib)',     'WhatsApp Servers',           'Meta / WhatsApp',        'HTTPS + WebSocket',     'QR Code Session',          'ثنائي / Bidirectional'],
    ['6', 'Flask Backend\n(ibulk_sms_service)','iBulk SMS API',            'iBulk SMS — Oman',       'HTTPS / REST',          'Username + Password',      'أحادي / Outbound only'],
    ['7', 'Flask Backend\n(pywebpush)',       'VAPID Push Servers',         'Google FCM / Apple APNs','HTTPS (Web Push W3C)',  'VAPID Private Key',        'أحادي / Outbound only'],
    ['8', 'Browser PWA\n(Service Worker)',    'Push Servers',               'Google FCM / Apple APNs','HTTPS + WebSocket',    'VAPID Public Key',         'أحادي / Inbound to browser'],
    ['9', 'Nginx',                            'Flask Backend',              'Self-hosted',            'HTTP (internal proxy)', 'Proxy Pass',               'ثنائي / Bidirectional'],
    ['10','Nginx',                            'Evolution API',              'Self-hosted',            'HTTP (internal proxy)', 'Proxy Pass',               'ثنائي / Bidirectional'],
]
add_table(doc, headers, rows, col_widths=[0.25, 1.2, 1.2, 1.3, 1.15, 1.2, 1.2])

# ── SECTION 3: Detailed per-connection ───────────────────────────────────────

add_heading(doc, '3. تفصيل كل اتصال / Detailed Connection Breakdown', 1)

# ── 3.1 ─────────────────────────────────────────────────────────────────────
add_heading(doc, '3.1  React Frontend  ↔  Flask Backend API', 2)
add_paragraph(doc,
    'الواجهة الأمامية المبنية بـ React تتواصل مع الـ Backend عبر HTTPS باستخدام Axios. '
    'كل طلب يحمل JWT Token في الـ Authorization Header للتحقق من هوية المستخدم.',
    size=10, rtl=True)
add_info_box(doc, [
    ('المزود / Provider',        'Self-hosted (خوادم المدرسة / VPS)'),
    ('البروتوكول',                'HTTPS + REST (JSON)'),
    ('مكتبة الاتصال',             'Axios 1.4.0'),
    ('المصادقة',                  'Authorization: Bearer <JWT_TOKEN>'),
    ('المنفذ الخارجي',            '443 (HTTPS via Nginx)'),
    ('المنفذ الداخلي',            '127.0.0.1:5000 (Flask)'),
    ('الحماية',                   'CORS + Rate Limiting + JWT Verification'),
    ('صيغة البيانات',             'JSON (Request & Response)'),
])

add_paragraph(doc, 'أهم نقاط النهاية / Key Endpoints:', bold=True, size=10)
ep_headers = ['الوظيفة / Function', 'Method', 'Endpoint']
ep_rows = [
    ['تسجيل الدخول',           'POST', '/api/auth/login'],
    ['تسجيل الحضور',           'POST', '/api/attendance/takes'],
    ['مسح QR الحافلة',         'POST', '/api/bus/scan'],
    ['إرسال إشعار',            'POST', '/api/notifications'],
    ['الاشتراك في Push',       'POST', '/api/notifications/subscribe'],
    ['طلب استلام طالب',        'POST', '/api/parent-pickup/request-pickup'],
    ['إرسال تقرير SMS',        'POST', '/api/attendance/send-daily-sms-reports'],
    ['إرسال WhatsApp',         'POST', '/api/auth/send-absence-notifications'],
]
add_table(doc, ep_headers, ep_rows, col_widths=[2.2, 0.8, 3.5], header_bg='2E5FA3')

# ── 3.2 ─────────────────────────────────────────────────────────────────────
add_heading(doc, '3.2  Flask Backend  ↔  MySQL Database', 2)
add_info_box(doc, [
    ('المزود / Provider',      'Self-hosted'),
    ('البروتوكول',              'TCP :3306'),
    ('ORM / مكتبة',             'SQLAlchemy 2.0.36 + PyMySQL 1.1.1'),
    ('المصادقة',                'Username + Password (environment variable)'),
    ('الشبكة',                  'داخلي فقط — 127.0.0.1:3306'),
    ('الحماية من SQL Injection','Parameterized Queries عبر SQLAlchemy ORM'),
    ('عزل البيانات',            'كل استعلام مُقيَّد بـ school_id من الـ JWT Token'),
])

# ── 3.3 ─────────────────────────────────────────────────────────────────────
add_heading(doc, '3.3  Flask Backend  ↔  Redis Cache', 2)
add_info_box(doc, [
    ('المزود / Provider',      'Self-hosted'),
    ('البروتوكول',              'RESP (Redis Protocol) / TCP :6379'),
    ('الاستخدام',               'Rate Limiting عبر Flask-Limiter'),
    ('المصادقة',                'Password (اختياري — يُنصح به)'),
    ('الشبكة',                  'داخلي فقط — 127.0.0.1:6379'),
    ('البيانات المخزنة',        'عدادات Rate Limiting فقط (لا بيانات مستخدمين)'),
    ('TTL',                    'مفاتيح تنتهي تلقائياً بعد نافذة الـ Rate Limit'),
])

# ── 3.4 ─────────────────────────────────────────────────────────────────────
add_heading(doc, '3.4  Flask Backend  →  Evolution API (WhatsApp)', 2)
add_paragraph(doc,
    'عند تسجيل غياب أو حدث على الحافلة، يرسل الـ Backend رسالة WhatsApp لولي الأمر عبر Evolution API '
    'المستضاف محلياً داخل Docker Container.',
    size=10, rtl=True)
add_info_box(doc, [
    ('المزود / Provider',      'Self-hosted — Evolution API (Open Source)'),
    ('البروتوكول',              'HTTP داخلي → REST API'),
    ('المصادقة',                'Header: apikey: <EVOLUTION_API_KEY>'),
    ('الاتجاه',                 'أحادي — Flask يرسل فقط'),
    ('الشبكة',                  'داخلي — localhost:8080 أو Docker bridge network'),
    ('Endpoint المستخدم',       'POST {api_url}/message/sendText/{instance_name}'),
    ('البيانات المرسلة',        '{ number: "968XXXXXXXX", text: "نص الرسالة" }'),
    ('البيانات المُعادة',        '{ key: { id: "..." }, status: "PENDING" }'),
    ('الإعداد لكل مدرسة',       'instance_name + api_key + api_url مستقلة لكل مدرسة'),
    ('التخزين',                  'الرسائل لا تُخزَّن في قاعدة بيانات التطبيق'),
])

# ── 3.5 ─────────────────────────────────────────────────────────────────────
add_heading(doc, '3.5  Evolution API  ↔  WhatsApp Servers (Meta)', 2)
add_info_box(doc, [
    ('المزود / Provider',      'Meta / WhatsApp'),
    ('البروتوكول',              'HTTPS + WebSocket (WhatsApp Web Protocol)'),
    ('مكتبة الاتصال',           'Baileys 7.0.0-rc.9 (مفتوح المصدر)'),
    ('المصادقة',                'QR Code Session — مسح مرة واحدة لكل مدرسة'),
    ('الاتجاه',                 'ثنائي (إرسال واستقبال)'),
    ('تخزين الجلسة',            'Docker Volume: evolution_instances'),
    ('البيانات المرسلة',        'رقم هاتف المستقبِل + نص الرسالة'),
    ('ملاحظة',                  'يعتمد على WhatsApp Web — غير رسمي من Meta'),
])

# ── 3.6 ─────────────────────────────────────────────────────────────────────
add_heading(doc, '3.6  Flask Backend  →  iBulk SMS API', 2)
add_info_box(doc, [
    ('المزود / Provider',      'iBulk SMS — مزود SMS عُماني محلي'),
    ('البروتوكول',              'HTTPS / REST'),
    ('المصادقة',                'Username + Password (مخزنة لكل مدرسة في MySQL)'),
    ('الاتجاه',                 'أحادي — Flask يرسل فقط'),
    ('البيانات المرسلة',        'رقم الهاتف + نص الرسالة + Sender ID'),
    ('البيانات المُعادة',        'حالة الإرسال + رصيد الحساب'),
    ('الإعداد لكل مدرسة',       'Credentials مستقلة لكل مدرسة'),
    ('الاستخدام',               'تقارير الحضور اليومية + إشعارات الغياب الجماعية'),
])

add_paragraph(doc, 'العمليات المتاحة / Available Operations:', bold=True, size=10)
add_table(doc, ['الدالة / Function', 'الوصف / Description'],
          [
              ['send_sms()', 'إرسال رسالة SMS لهاتف واحد'],
              ['send_bulk_sms()', 'إرسال SMS لمجموعة من الأرقام دفعة واحدة'],
              ['check_balance()', 'التحقق من الرصيد المتبقي في الحساب'],
          ], col_widths=[2.0, 4.5], header_bg='2E5FA3')

# ── 3.7 ─────────────────────────────────────────────────────────────────────
add_heading(doc, '3.7  Flask Backend  →  VAPID Push Servers (Google FCM / Apple APNs)', 2)
add_info_box(doc, [
    ('المزود / Provider',      'Google Firebase Cloud Messaging (FCM) + Apple APNs'),
    ('البروتوكول',              'HTTPS + W3C Web Push Standard'),
    ('مكتبة الاتصال',           'pywebpush 1.14.1'),
    ('المصادقة',                'VAPID Private Key (مخزن في .env فقط — لا يُرسَل للعميل)'),
    ('الاتجاه',                 'أحادي — Flask → Push Server → Browser'),
    ('تشفير الحمولة',           'End-to-end encryption باستخدام مفتاح الاشتراك (p256dh)'),
    ('التسليم',                  'يُسلَّم حتى لو كان التطبيق مغلقاً تماماً'),
    ('البيانات المرسلة',        'عنوان الإشعار + النص + رابط اختياري'),
])

# ── 3.8 ─────────────────────────────────────────────────────────────────────
add_heading(doc, '3.8  Browser PWA (Service Worker)  ↔  Push Servers', 2)
add_info_box(doc, [
    ('المزود / Provider',      'Google FCM / Apple APNs'),
    ('البروتوكول',              'HTTPS + PushManager API (W3C)'),
    ('المصادقة',                'VAPID Public Key (آمن للمشاركة)'),
    ('الاتجاه',                 'أحادي — Push Server → Browser'),
    ('الاشتراك',                'navigator.serviceWorker + PushManager.subscribe()'),
    ('التشغيل بدون التطبيق',    'Service Worker يعمل في الخلفية ويعرض الإشعار'),
    ('تخزين الـ Endpoint',      'push_subscriptions table في MySQL'),
])

# ── SECTION 4: Data Flow Per Event ───────────────────────────────────────────

add_heading(doc, '4. مسار البيانات لكل حدث / Data Flow Per Event', 1)

# event 1
add_heading(doc, '4.1  حدث: تسجيل غياب طالب', 2)
add_code_block(doc, """المعلم (Browser)
    │
    1──► POST /api/attendance/takes           [HTTPS → Nginx → Flask]
    │
    2──► Flask يكتب في MySQL                 [TCP داخلي :3306]
    │
    3──► Flask → Evolution API               [HTTP داخلي :8080]
    │         └──► Evolution → WhatsApp Servers [HTTPS + WebSocket]
    │                   └──► هاتف ولي الأمر  📱
    │
    4──► Flask → VAPID Push Servers          [HTTPS / Web Push]
    │         └──► متصفح الطالب / ولي الأمر  🔔
    │
    5──► استجابة للمعلم ✅                   [JSON / HTTPS]""")

# event 2
add_heading(doc, '4.2  حدث: مسح QR الحافلة', 2)
add_code_block(doc, """السائق (Mobile Browser)
    │
    1──► كاميرا تقرأ QR Code الطالب
    │
    2──► POST /api/bus/scan                   [HTTPS → Flask]
    │
    3──► Flask يكتب BusScan في MySQL          [TCP داخلي :3306]
    │
    4──► Flask → VAPID Push Servers → FCM     [HTTPS]
    │         └──► ولي الأمر: "صعد ابنك الحافلة"  🔔
    │
    5──► استجابة للسائق: اسم الطالب + وقت المسح ✅""")

# event 3
add_heading(doc, '4.3  حدث: إرسال تقرير SMS يومي', 2)
add_code_block(doc, """الإدارة (Browser)
    │
    1──► POST /api/attendance/send-daily-sms-reports   [HTTPS → Flask]
    │
    2──► Flask يجلب بيانات الحضور من MySQL            [TCP داخلي :3306]
    │
    3──► Flask → iBulk SMS API                         [HTTPS]
    │         └──► iBulk → شبكات الاتصال العُمانية
    │                   └──► ولي الأمر يستقبل SMS  📩
    │
    4──► استجابة: عدد الرسائل + حالة الرصيد ✅""")

# event 4
add_heading(doc, '4.4  حدث: تسجيل الدخول', 2)
add_code_block(doc, """المستخدم (Browser)
    │
    1──► POST /api/auth/login  { username, password }  [HTTPS]
    │
    2──► Flask يتحقق من Redis: هل تجاوز Rate Limit؟  [TCP داخلي :6379]
    │
    3──► Flask يجلب المستخدم من MySQL                 [TCP داخلي :3306]
         └──► يتحقق من PBKDF2-SHA256 password hash
    │
    4──► Flask يسجل في action_logs (IP + وقت)        [MySQL]
    │
    5──► استجابة: { access_token, refresh_token, user }  ✅""")

# ── SECTION 5: Providers Summary ─────────────────────────────────────────────

add_heading(doc, '5. ملخص المزودين / Providers Summary', 1)

p_headers = ['المزود\nProvider', 'الخدمة\nService', 'الموقع\nLocation',
             'نوع الاستضافة\nHosting', 'البروتوكول\nProtocol']
p_rows = [
    ['Self-hosted',            'Flask API',           'السيرفر الخاص',    'On-premise / VPS',     'HTTP داخلي'],
    ['Self-hosted',            'MySQL Database',      'السيرفر الخاص',    'On-premise / VPS',     'TCP :3306'],
    ['Self-hosted',            'Redis',               'السيرفر الخاص',    'On-premise / VPS',     'TCP :6379'],
    ['Self-hosted (Docker)',   'Evolution API',       'السيرفر الخاص',    'Docker Container',     'HTTP :8080'],
    ['Meta / WhatsApp',        'WhatsApp Servers',    'خوادم Meta العالمية','Cloud (External)',   'HTTPS + WebSocket'],
    ['iBulk SMS — Oman',       'SMS Gateway',         'عُمان',            'Cloud (External)',     'HTTPS / REST'],
    ['Google',                 'Firebase FCM',        'خوادم Google',     'Cloud (External)',     'HTTPS / Web Push'],
    ['Apple',                  'APNs',                'خوادم Apple',      'Cloud (External)',     'HTTPS / Web Push'],
]
add_table(doc, p_headers, p_rows,
          col_widths=[1.4, 1.3, 1.4, 1.3, 1.3])

# ── SECTION 6: Ports Map ─────────────────────────────────────────────────────

add_heading(doc, '6. خريطة المنافذ / Ports & Network Map', 1)

add_paragraph(doc, 'المنافذ المفتوحة للخارج فقط / Externally Exposed Ports:', bold=True, size=10)
add_table(doc, ['المنفذ / Port', 'الخدمة / Service', 'ملاحظة / Note'],
          [
              ['443', 'HTTPS',  'كل حركة البيانات — All application traffic'],
              ['80',  'HTTP',   'يُعاد توجيهه فوراً لـ 443 / Redirect to HTTPS only'],
              ['22',  'SSH',    'للإدارة فقط — مفتاح خاص / Key-based auth only'],
          ], col_widths=[1.0, 1.5, 4.0], header_bg='2E5FA3')

add_paragraph(doc, 'المنافذ الداخلية فقط (محظورة من الخارج) / Internal-Only Ports:', bold=True, size=10)
add_table(doc, ['المنفذ / Port', 'الخدمة / Service', 'الربط / Binding'],
          [
              ['5000', 'Flask Backend API',    '127.0.0.1:5000'],
              ['8080', 'Evolution API',         '127.0.0.1:8080'],
              ['4000', 'Evolution Manager UI',  '127.0.0.1:4000'],
              ['3306', 'MySQL Database',         '127.0.0.1:3306'],
              ['6379', 'Redis Cache',            '127.0.0.1:6379'],
          ], col_widths=[1.0, 2.5, 3.0], header_bg='555555')

# ── save ──────────────────────────────────────────────────────────────────────

output_path = r'c:\Users\User\Desktop\PathToDiv\tatubujs\docs\API_CONNECTIONS_MAP.docx'
doc.save(output_path)
print(f'[OK] Word document saved: {output_path}')
